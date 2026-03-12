import io
from django.http import FileResponse
from rest_framework import viewsets
from rest_framework.decorators import action
from rest_framework.permissions import IsAuthenticated
from rest_framework.response import Response

from .models import LessonPlan
from .serializers import LessonPlanSerializer


TOOL_TYPE_LABELS = dict(LessonPlan.TOOL_TYPE_CHOICES)


def _build_export_docx(lesson):
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    content = lesson.content or {}
    tool_type = getattr(lesson, 'tool_type', 'lesson_plan') or 'lesson_plan'

    title = content.get('Тема_урока') or lesson.title
    p = doc.add_heading(title, level=0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    label = TOOL_TYPE_LABELS.get(tool_type, tool_type)
    doc.add_paragraph(f"Тип: {label}")

    div = content.get('Раздел')
    if div:
        doc.add_paragraph(f"Раздел: {div}")

    goals_learn = content.get('Цели_обучения') or []
    if goals_learn:
        doc.add_heading('Цели обучения', level=1)
        for g in goals_learn:
            doc.add_paragraph(str(g), style='List Bullet')

    goals_lesson = content.get('Цели_урока') or []
    if goals_lesson:
        doc.add_heading('Цели урока', level=1)
        for g in goals_lesson:
            doc.add_paragraph(str(g), style='List Bullet')

    if tool_type == 'lesson_plan':
        _export_lesson_plan(doc, content)
    elif tool_type == 'sor_soch':
        _export_assessment(doc, content)
    elif tool_type == 'formative':
        _export_formative(doc, content)
    elif tool_type == 'structural':
        _export_structural(doc, content)
    elif tool_type == 'presentation':
        _export_presentation(doc, content)
    elif tool_type == 'feedback':
        _export_feedback(doc, content)
    elif tool_type == 'organizer':
        _export_organizer(doc, content)
    elif tool_type == 'game':
        _export_game(doc, content)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def _export_lesson_plan(doc, content):
    hod = content.get('Ход_урока') or []
    if hod:
        doc.add_heading('Ход урока', level=1)
        table = doc.add_table(rows=1 + len(hod), cols=6)
        table.style = 'Table Grid'
        headers = ['Этап / Время', 'Модель урока', 'Действия педагога', 'Действия ученика', 'Ресурсы', 'Оценивание']
        for i, text in enumerate(headers):
            table.rows[0].cells[i].text = text
        for row_idx, step in enumerate(hod, start=1):
            row = table.rows[row_idx].cells
            row[0].text = f"{step.get('Этап', '')} {step.get('Время', '')}".strip()
            row[1].text = str(step.get('Модель урока', ''))
            row[2].text = str(step.get('Действия педагога', ''))
            row[3].text = str(step.get('Действия ученика', ''))
            row[4].text = str(step.get('Ресурсы', ''))
            row[5].text = str(step.get('Оценивание', ''))


def _export_assessment(doc, content):
    tasks = content.get('tasks') or []
    if not tasks:
        return
    doc.add_heading('Задания СОР/СОЧ', level=1)
    for i, task in enumerate(tasks, start=1):
        q_type = task.get('question_type', '')
        q_level = task.get('question_level', '')
        doc.add_paragraph(f"{i}. [{q_type}] [{q_level}]", style='Heading 2')
        doc.add_paragraph(task.get('question', ''))
        for opt in (task.get('options') or []):
            if isinstance(opt, dict):
                mark = ' ✓' if (opt.get('is_correct') is True or opt.get('is_correct') == 'true') else ''
                doc.add_paragraph(f"  • {opt.get('option', '')}{mark}", style='List Bullet')
            else:
                doc.add_paragraph(f"  • {opt}", style='List Bullet')
        if task.get('correct_answer_explanation'):
            doc.add_paragraph(f"Обоснование: {task['correct_answer_explanation']}")
        if task.get('resource'):
            doc.add_paragraph(f"Ресурс: {task['resource']}")


def _export_formative(doc, content):
    exit_tickets = content.get('exit_ticket') or []
    if exit_tickets:
        doc.add_heading('Выходной билет (Exit Ticket)', level=1)
        for i, item in enumerate(exit_tickets, start=1):
            doc.add_paragraph(f"{i}. {item.get('question', '')}")
            doc.add_paragraph(f"   Ожидаемый ответ: {item.get('expected_answer', '')}")

    polls = content.get('quick_poll') or []
    if polls:
        doc.add_heading('Быстрый опрос', level=1)
        for i, item in enumerate(polls, start=1):
            correct = 'Верно' if item.get('correct') else 'Неверно'
            doc.add_paragraph(f"{i}. {item.get('statement', '')} — {correct}")
            if item.get('explanation'):
                doc.add_paragraph(f"   Пояснение: {item['explanation']}")

    sa = content.get('self_assessment') or []
    if sa:
        doc.add_heading('Лист самооценки', level=1)
        table = doc.add_table(rows=1 + len(sa), cols=4)
        table.style = 'Table Grid'
        table.rows[0].cells[0].text = 'Критерий'
        table.rows[0].cells[1].text = 'Могу'
        table.rows[0].cells[2].text = 'Частично'
        table.rows[0].cells[3].text = 'Нужна помощь'
        for idx, item in enumerate(sa, start=1):
            table.rows[idx].cells[0].text = str(item.get('criteria', ''))
            table.rows[idx].cells[1].text = str(item.get('descriptor_can', ''))
            table.rows[idx].cells[2].text = str(item.get('descriptor_partial', ''))
            table.rows[idx].cells[3].text = str(item.get('descriptor_need_help', ''))

    obs = content.get('observation_sheet') or []
    if obs:
        doc.add_heading('Наблюдательный лист', level=1)
        for item in obs:
            doc.add_paragraph(f"Критерий: {item.get('criteria', '')}")
            doc.add_paragraph(f"Показатели: {item.get('indicators', '')}")

    kwl = content.get('kwl')
    if kwl:
        doc.add_heading('KWL таблица', level=1)
        table = doc.add_table(rows=2, cols=3)
        table.style = 'Table Grid'
        table.rows[0].cells[0].text = 'Знаю (K)'
        table.rows[0].cells[1].text = 'Хочу узнать (W)'
        table.rows[0].cells[2].text = 'Узнал (L)'
        table.rows[1].cells[0].text = '\n'.join(kwl.get('know_prompts') or [])
        table.rows[1].cells[1].text = '\n'.join(kwl.get('want_prompts') or [])
        table.rows[1].cells[2].text = '\n'.join(kwl.get('learned_prompts') or [])


def _export_structural(doc, content):
    levels = content.get('levels') or []
    if not levels:
        return
    doc.add_heading('Структурные задания', level=1)
    for level_data in levels:
        lvl = level_data.get('level', '')
        name = level_data.get('level_name', '')
        doc.add_heading(f"Уровень {lvl} — {name}", level=2)
        for task in (level_data.get('tasks') or []):
            doc.add_paragraph(f"Задание {task.get('task_number', '')}: {task.get('task_text', '')}", style='List Number')
            if task.get('instruction'):
                doc.add_paragraph(f"Инструкция: {task['instruction']}")
            for desc in (task.get('descriptors') or []):
                doc.add_paragraph(f"  • {desc.get('criterion', '')} (макс. {desc.get('max_score', '')} б.)", style='List Bullet')
            doc.add_paragraph(f"Макс. балл: {task.get('max_score', '')}, Время: {task.get('time_minutes', '')} мин.")


def _export_presentation(doc, content):
    slides = content.get('Презентация') or []
    if not slides:
        return
    doc.add_heading('Презентация', level=1)
    for slide in slides:
        doc.add_heading(f"{slide.get('Слайд', '')}: {slide.get('Заголовок', '')}", level=2)
        if slide.get('Контекст'):
            doc.add_paragraph(slide['Контекст'])
        if slide.get('Иллюстрация'):
            doc.add_paragraph(f"Иллюстрация: {slide['Иллюстрация']}")


def _export_feedback(doc, content):
    fb = content.get('Обратная_связь') or content.get('Обратная связь') or []
    if fb:
        doc.add_heading('Обратная связь', level=1)
        for item in fb:
            doc.add_heading(item.get('Принцип', ''), level=2)
            doc.add_paragraph(item.get('Комментарий', ''))

    questions = content.get('Вопросы_для_рефлексии') or []
    if questions:
        doc.add_heading('Вопросы для рефлексии', level=1)
        for q in questions:
            doc.add_paragraph(str(q), style='List Bullet')

    recs = content.get('Рекомендации') or []
    if recs:
        doc.add_heading('Рекомендации', level=1)
        for r in recs:
            doc.add_paragraph(str(r), style='List Bullet')


def _export_organizer(doc, content):
    title = content.get('title', content.get('topic', 'Органайзер'))
    doc.add_heading(title, level=1)
    doc.add_paragraph(f"Тип: {content.get('type', '')}")
    doc.add_paragraph(f"Тема: {content.get('topic', '')}")
    doc.add_paragraph()
    doc.add_paragraph(
        "Этот органайзер является интерактивным HTML-документом. "
        "Откройте его на платформе EduDesk для редактирования и используйте кнопку «Скачать как Word» "
        "для получения полноценного Word-файла с вашими правками.",
        style='Intense Quote'
    )


def _export_game(doc, content):
    title = content.get('title', content.get('topic', 'Игра'))
    doc.add_heading(title, level=1)
    doc.add_paragraph(f"Тип игры: {content.get('type', '')}")
    doc.add_paragraph(f"Тема: {content.get('topic', '')}")
    doc.add_paragraph()
    doc.add_paragraph(
        "Эта игра является интерактивной HTML/JS страницей. "
        "Откройте её в браузере на платформе EduDesk для полноценного игрового процесса.",
        style='Intense Quote'
    )


class LessonPlanViewSet(viewsets.ModelViewSet):
    serializer_class = LessonPlanSerializer
    permission_classes = [IsAuthenticated]

    def get_queryset(self):
        qs = LessonPlan.objects.filter(user=self.request.user).order_by('-created_at')
        tool_type = self.request.query_params.get('tool_type')
        if tool_type:
            qs = qs.filter(tool_type=tool_type)
        return qs

    def perform_create(self, serializer):
        serializer.save(user=self.request.user)

    @action(detail=True, methods=['get'], url_path='export')
    def export(self, request, pk=None):
        lesson = self.get_object()
        fmt = (request.query_params.get('format') or 'docx').lower()
        if fmt != 'docx':
            return Response({'error': 'Поддерживается только format=docx'}, status=400)
        try:
            buffer = _build_export_docx(lesson)
        except ImportError:
            return Response({'error': 'Установите python-docx: pip install python-docx'}, status=500)
        filename = f"{lesson.tool_type}_{lesson.id}_{lesson.title[:30].replace(' ', '_')}.docx"
        return FileResponse(
            buffer, as_attachment=True, filename=filename,
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        )
